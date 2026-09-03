"""DOCX signature-block extraction — tables (incl. nested) and text-boxes.

Construction letters put the sign-off (name / role / company) in a nested
table or a Word text-box (w:txbxContent), not in body paragraphs.
python-docx `.paragraphs` + top-level `.tables` / `cell.text` drop that
text, so the indexer never saw "Barry Muir". These fixtures prove the
production extractor used by doc_index (`extract_document_text`) now
returns that text, without dumping XML tags or depending on Drive/prod.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.core.doc_index import extract_document_text

BODY_PHRASE = "UBCC Concrete Batching Plant at Wadi Safar"
BODY_DATE = "March 2024"
BODY_WORKS = "in-situ"

SIGNATORY = "Barry Muir"
ROLE = "Engineer's Representative"
COMPANY = "CH2M Saudi Limited"

# VML text-box — the same w:txbxContent Word uses for floating signature
# frames. python-docx has no high-level API for this; we inject the XML.
_TXBX_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml"'
)
_SIGNATURE_TEXTBOX_XML = (
    f"<w:r {_TXBX_NS}>"
    "<w:pict>"
    '<v:shape style="width:220pt;height:80pt">'
    "<v:textbox>"
    "<w:txbxContent>"
    f"<w:p><w:r><w:t>{ROLE}</w:t></w:r></w:p>"
    f"<w:p><w:r><w:t>{COMPANY}</w:t></w:r></w:p>"
    "</w:txbxContent>"
    "</v:textbox>"
    "</v:shape>"
    "</w:pict>"
    "</w:r>"
)


def _write_letter_docx(path: Path) -> Path:
    """Letter whose body lacks the signatory; name/role/company are only
    in a nested table and a text-box."""
    from lxml import etree
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        f"Letter to AICC on Completion and transfer of responsibility — {BODY_PHRASE}."
    )
    doc.add_paragraph(f"{BODY_DATE}. Works include {BODY_WORKS} concrete.")

    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = SIGNATORY

    host = doc.add_paragraph()
    host._p.append(etree.fromstring(_SIGNATURE_TEXTBOX_XML))

    doc.save(path)
    return path


def _legacy_paragraphs_and_toplevel_tables(path: Path) -> str:
    """The extractor as it was: body paragraphs + top-level cell.text only.

    Nested w:tbl is invisible to cell.text; w:txbxContent is invisible to
    both paragraphs and tables. Used to pin the before-fix miss.
    """
    from docx import Document

    document = Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for tbl in getattr(document, "tables", None) or []:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@pytest.fixture
def letter_docx(tmp_path: Path) -> Path:
    return _write_letter_docx(tmp_path / "ubcc_completion_letter.docx")


def test_legacy_path_drops_nested_table_and_textbox_signature(letter_docx: Path):
    """Before-fix behaviour: signatory lives only in nested table / text-box."""
    legacy = _legacy_paragraphs_and_toplevel_tables(letter_docx)
    assert BODY_PHRASE in legacy
    assert SIGNATORY not in legacy
    assert ROLE not in legacy
    assert COMPANY not in legacy


def test_extract_document_text_indexes_table_and_textbox_signature(
    letter_docx: Path, monkeypatch: pytest.MonkeyPatch
):
    """Production doc_index path must return all three signature strings."""
    monkeypatch.delenv("DATA_ENCRYPTION_KEY", raising=False)
    text = extract_document_text(str(letter_docx), letter_docx.name)

    # Body paragraphs still extracted.
    assert BODY_PHRASE in text
    assert BODY_DATE in text
    assert BODY_WORKS in text

    # Nested-table name + text-box role/company — the dropped signature.
    assert SIGNATORY in text
    assert ROLE in text
    assert COMPANY in text

    # Plaintext only: no OOXML leak into the index.
    assert "txbxContent" not in text
    assert "txBody" not in text
    assert "<w:" not in text
    assert "xmlns:" not in text
