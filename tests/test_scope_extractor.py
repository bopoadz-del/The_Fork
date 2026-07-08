"""Tests for the roadmap-named Scope Extractor shim.

The real document extraction is tested elsewhere; this file only proves that
``app.blocks.scope_extractor.ScopeExtractorBlock`` delegates to the document
engine and returns the expected schedule-scope shape.
"""

import pytest

from app.blocks.scope_extractor import ScopeExtractorBlock


@pytest.fixture
def block():
    return ScopeExtractorBlock()


@pytest.mark.asyncio
async def test_scope_extractor_returns_error_when_no_input(block):
    result = await block.process({})
    # The underlying document engine returns an error when no files/text.
    assert result["status"] == "error"


@pytest.mark.asyncio
async def test_scope_extractor_focuses_output_shape(block, tmp_path):
    from docx import Document

    doc_path = tmp_path / "BOD.docx"
    doc = Document()
    doc.add_paragraph("Data center with 200 day generator lead time.")
    doc.add_paragraph("Target: operational by December 2027.")
    doc.save(str(doc_path))

    result = await block.process({"file_path": str(doc_path)})
    # Document engine may succeed or error depending on config; just ensure shape
    # is schedule-scope when it succeeds.
    assert "status" in result
    if result["status"] == "success":
        assert "schedule_targets" in result
        assert "equipment_lead_times" in result
        assert "target_milestones" in result


def test_scope_extractor_block_metadata():
    assert ScopeExtractorBlock.name == "scope_extractor"
    assert "scope" in ScopeExtractorBlock.tags
