"""The schedule renderers must put the DATA in the file, not just make a file.

`_render_xlsx` and `_render_docx` were both at ~3% of body executed: the export
endpoints were exercised, but only far enough to confirm a 200 and a non-zero
byte count. A renderer that writes a correct header row and then drops every
activity produces a perfectly valid, perfectly empty workbook -- the exact
wrap-around green this codebase refuses elsewhere.

So these tests open the artefact and read the activities back out.
"""
from __future__ import annotations

import os

import pytest

from app.routers.exports import _render_docx, _render_xlsx

ACTIVITIES = [
    {"id": "A10", "wbs": "1.1", "name": "Excavation", "duration": 10,
     "start": "2026-01-05", "finish": "2026-01-16", "predecessors": [],
     "resources": "Excavator crew"},
    {"id": "A20", "wbs": "1.2", "name": "Raft pour", "duration": 5,
     "start": "2026-01-19", "finish": "2026-01-23", "predecessors": ["A10"],
     "resources": "Concrete gang"},
]


@pytest.fixture
def cleanup():
    made: list[str] = []
    yield made
    for path in made:
        try:
            os.unlink(path)
        except OSError:
            pass


def _xlsx_text(path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(path)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if c.value is not None:
                    out.append(str(c.value))
    return "\n".join(out)


def _docx_text(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ── xlsx ──────────────────────────────────────────────────────────────────

def test_xlsx_contains_every_activity_not_just_the_header(cleanup):
    path = _render_xlsx(ACTIVITIES, "Tower A", "2026-01-05", None)
    cleanup.append(path)
    assert os.path.getsize(path) > 0

    body = _xlsx_text(path)
    for act in ACTIVITIES:
        assert act["id"] in body, f"activity {act['id']} missing from the workbook"
        assert act["name"] in body, f"name {act['name']!r} missing"
    assert "Tower A" in body
    assert "2026-01-05" in body


def test_xlsx_activity_count_in_the_banner_matches_the_rows(cleanup):
    """The banner reports a count. If it disagrees with what was written, the
    reader trusts the banner and never notices rows were dropped."""
    path = _render_xlsx(ACTIVITIES, "Tower A", "2026-01-05", None)
    cleanup.append(path)
    body = _xlsx_text(path)
    assert f"Activities: {len(ACTIVITIES)}" in body


def test_xlsx_with_no_activities_still_produces_a_readable_file(cleanup):
    """An empty schedule is a legitimate answer; a corrupt file is not."""
    path = _render_xlsx([], "Empty", None, None)
    cleanup.append(path)
    body = _xlsx_text(path)
    assert "Empty" in body
    assert "Activities: 0" in body


def test_xlsx_notes_reach_the_workbook(cleanup):
    path = _render_xlsx(ACTIVITIES, "Tower A", "2026-01-05",
                        "Sequence assumes early piling release.")
    cleanup.append(path)
    assert "early piling release" in _xlsx_text(path)


# ── docx ──────────────────────────────────────────────────────────────────

def test_docx_contains_every_activity(cleanup):
    path = _render_docx(ACTIVITIES, "Tower A", "2026-01-05", None)
    cleanup.append(path)
    assert os.path.getsize(path) > 0

    body = _docx_text(path)
    for act in ACTIVITIES:
        assert act["id"] in body, f"activity {act['id']} missing from the docx"
        assert act["name"] in body
    assert "Tower A" in body


def test_docx_with_no_activities_is_still_a_valid_document(cleanup):
    path = _render_docx([], "Empty", None, None)
    cleanup.append(path)
    body = _docx_text(path)
    assert "Empty" in body


def test_docx_notes_reach_the_document(cleanup):
    path = _render_docx(ACTIVITIES, "Tower A", "2026-01-05",
                        "Float consumed by procurement.")
    cleanup.append(path)
    assert "Float consumed by procurement." in _docx_text(path)


def test_renderers_write_distinct_files_per_call(cleanup):
    """Concurrent exports must not collide on one temp path and hand two users
    the same file."""
    a = _render_xlsx(ACTIVITIES, "One", None, None)
    b = _render_xlsx(ACTIVITIES, "Two", None, None)
    cleanup.extend([a, b])
    assert a != b
    assert "One" in _xlsx_text(a)
    assert "Two" in _xlsx_text(b)
