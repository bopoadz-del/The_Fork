"""Schedule + report export endpoints — xlsx / docx / pdf.

Takes a list of activity dicts (the shape `generate_wbs` returns in its
`activities` field) and renders them to a file the user can download. The
schedule export wraps openpyxl directly (works without the existing
`write_schedule_excel` helper because that helper expects a CPMOutput
dataclass instance — these endpoints accept the plain dicts that the
agent runtime emits).

DOCX rendering uses python-docx; PDF rendering uses reportlab if
installed and otherwise falls back to a docx-as-pdf fallback (which on
this image is itself unavailable, so PDF returns 503 when reportlab is
absent).
"""
from __future__ import annotations

import io
import logging
import os
import tempfile
import uuid
from typing import Any, Dict, List, Optional

from datetime import datetime, timezone

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from app.dependencies import require_user
from app.core import projects as projects_store
from app.core import agent_memory
from app.core import doc_index, file_crypto

router = APIRouter()

logger = logging.getLogger(__name__)

# Mirror the upload path's storage root so generated workbooks persist next to
# uploaded documents and feed the same RAG indexer (see app/routers/projects.py).
DATA_DIR = os.getenv("DATA_DIR", "./data")


class ScheduleExportRequest(BaseModel):
    """Inline activity list. Each activity is a dict with at least
    `id`, `name`, `duration_days`, `early_start_day`, `early_finish_day`,
    `total_float_days`, `critical`, `wbs_phase`, `predecessors`.
    """
    activities: List[Dict[str, Any]] = Field(..., description="generate_wbs-style activity rows")
    project_name: Optional[str] = None
    start_date: Optional[str] = None  # ISO YYYY-MM-DD
    end_date: Optional[str] = None
    notes: Optional[str] = None


class CostBoqExportRequest(BaseModel):
    """Generate a formula-linked cost BOQ. Either pass structured ``categories``
    or a ``document_id`` of an uploaded priced BOQ to derive them from."""
    title: Optional[str] = None
    project_name: Optional[str] = None
    location: Optional[str] = None
    currency: str = "SAR"
    date: Optional[str] = None
    categories: Optional[List[Dict[str, Any]]] = None
    document_id: Optional[str] = None
    # When true (default) the generated workbook is also persisted back into the
    # project and queued for eager indexing so chat can answer from it. Set
    # false to get the download only, without touching the RAG corpus.
    ingest: bool = True


class PriceBoqRequest(BaseModel):
    """Price an UNPRICED BOQ from the typed rate-card.

    ``document_id`` is an uploaded digital/xlsx BOQ whose rates are blank/0.
    ``asset_type`` selects the rate-card (e.g. "Buildings/Towers"); ``currency``
    defaults to that asset's sole/first currency. The generated workbook is
    persisted + eager-indexed unless ``ingest`` is false."""
    document_id: str
    asset_type: str
    currency: Optional[str] = None
    project_name: Optional[str] = None
    ingest: bool = True


class CostScheduleExportRequest(BaseModel):
    """Cost-loaded L2 schedule. Activities: {id, wbs, name, duration,
    predecessors:[id], cost, manpower}."""
    project_name: Optional[str] = None
    currency: str = "SAR"
    activities: List[Dict[str, Any]] = Field(..., description="CPM activities with cost + manpower")


class EvmExportRequest(BaseModel):
    """EVM workbook. Periods: {period, pv, ev, ac}; bac = budget at completion."""
    project_name: Optional[str] = None
    currency: str = "SAR"
    bac: float = 0
    periods: List[Dict[str, Any]] = Field(..., description="PV/EV/AC per period")


def _check_owner(project_id: str, user_id: str) -> Dict[str, Any]:
    proj = projects_store.get_project(project_id, user_id=user_id)
    if not proj:
        raise HTTPException(404, f"Project '{project_id}' not found")
    return proj


def _categories_from_document(project_id: str, document_id: str) -> List[Dict[str, Any]]:
    """Derive cost-BOQ categories from an uploaded priced BOQ via boq_processor
    (groups its line items by section). Raises 4xx if no priced items parse."""
    doc = projects_store.get_document(document_id)
    if not doc or not doc.get("file_path"):
        raise HTTPException(404, "document not found")
    from app.blocks.boq_processor import BOQProcessorBlock
    from app.core.doc_index import _run_sync
    res = _run_sync(BOQProcessorBlock().process(
        {"file_path": doc["file_path"], "project_id": project_id}))
    if res.get("status") != "success" or not res.get("line_items"):
        raise HTTPException(
            422, "could not extract priced line items from that document — a "
                 "digital/xlsx BOQ is required (scanned PDFs won't parse).")
    by_section: Dict[str, List[Dict[str, Any]]] = {}
    for i, it in enumerate(res["line_items"], 1):
        sec = it.get("section") or "General"
        by_section.setdefault(sec, []).append({
            "item_no": it.get("item_key") or str(i),
            "description": it.get("description") or "",
            "unit": it.get("unit") or "",
            "qty": it.get("quantity") or 0,
            "rate": it.get("unit_cost") or 0,
        })
    return [{"name": s, "items": items} for s, items in by_section.items()]


def _render_xlsx(activities: List[Dict[str, Any]], project_name: str,
                 start_date: Optional[str], notes: Optional[str]) -> str:
    """Write activities to a temp xlsx and return its path."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule"

    # Header banner
    ws["A1"] = project_name or "Schedule"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"Start: {start_date or ''}    Activities: {len(activities)}"
    ws["A2"].font = Font(italic=True)

    # Column headers
    headers = ["ID", "Phase", "Name", "Duration (days)", "ES day", "EF day",
               "Total Float", "Critical", "Predecessors", "Resources"]
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=ci, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F3864")
        cell.alignment = Alignment(horizontal="center")

    crit_fill = PatternFill("solid", fgColor="F8CBAD")
    for ri, a in enumerate(activities, start=5):
        critical = bool(a.get("critical"))
        row = [
            a.get("id") or a.get("code") or "",
            a.get("wbs_phase") or a.get("phase") or "",
            a.get("name") or "",
            a.get("duration_days") or a.get("duration") or 0,
            a.get("early_start_day") if a.get("early_start_day") is not None else "",
            a.get("early_finish_day") if a.get("early_finish_day") is not None else "",
            a.get("total_float_days") if a.get("total_float_days") is not None else "",
            "YES" if critical else "",
            ", ".join(a.get("predecessors") or []),
            ", ".join(a.get("resources") or []),
        ]
        for ci, v in enumerate(row, start=1):
            cell = ws.cell(row=ri, column=ci, value=v)
            if critical:
                cell.fill = crit_fill

    # Column widths
    widths = [10, 22, 40, 14, 10, 10, 12, 10, 28, 22]
    for ci, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + ci)].width = w

    if notes:
        notes_row = 5 + len(activities) + 2
        ws.cell(row=notes_row, column=1, value="Notes:").font = Font(bold=True)
        ws.cell(row=notes_row + 1, column=1, value=notes)

    fd, path = tempfile.mkstemp(prefix="schedule_", suffix=".xlsx")
    os.close(fd)
    wb.save(path)
    return path


def _render_docx(activities: List[Dict[str, Any]], project_name: str,
                 start_date: Optional[str], notes: Optional[str]) -> str:
    """Write activities to a temp .docx schedule report and return its path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()
    doc.add_heading(project_name or "Schedule", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Start: {start_date or '—'}    Activities: {len(activities)}").italic = True

    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for ci, h in enumerate(
        ["ID", "Phase", "Name", "Dur.", "ES", "EF", "Float", "Crit."]
    ):
        hdr[ci].text = h
        for p in hdr[ci].paragraphs:
            for r in p.runs:
                r.bold = True

    for a in activities:
        row = table.add_row().cells
        row[0].text = str(a.get("id") or a.get("code") or "")
        row[1].text = str(a.get("wbs_phase") or a.get("phase") or "")
        row[2].text = str(a.get("name") or "")
        row[3].text = str(a.get("duration_days") or a.get("duration") or "")
        row[4].text = str(a.get("early_start_day") if a.get("early_start_day") is not None else "")
        row[5].text = str(a.get("early_finish_day") if a.get("early_finish_day") is not None else "")
        row[6].text = str(a.get("total_float_days") if a.get("total_float_days") is not None else "")
        row[7].text = "YES" if a.get("critical") else ""

    if notes:
        doc.add_paragraph()
        nh = doc.add_paragraph()
        nh.add_run("Notes").bold = True
        doc.add_paragraph(notes)

    fd, path = tempfile.mkstemp(prefix="schedule_", suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


def _render_pdf(activities: List[Dict[str, Any]], project_name: str,
                start_date: Optional[str], notes: Optional[str]) -> str:
    """Write activities to a temp PDF using reportlab. Raises 503 if absent."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="PDF export unavailable — reportlab is not installed in this image. "
                   "Uncomment reportlab in requirements.txt and redeploy to enable.",
        )

    fd, path = tempfile.mkstemp(prefix="schedule_", suffix=".pdf")
    os.close(fd)
    pdf = SimpleDocTemplate(path, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = [
        Paragraph(project_name or "Schedule", styles["Title"]),
        Paragraph(
            f"Start: {start_date or '—'}    Activities: {len(activities)}",
            styles["Italic"],
        ),
        Spacer(1, 12),
    ]
    data = [["ID", "Phase", "Name", "Dur.", "ES", "EF", "Float", "Crit."]]
    for a in activities:
        data.append([
            str(a.get("id") or a.get("code") or ""),
            str(a.get("wbs_phase") or a.get("phase") or ""),
            (str(a.get("name") or ""))[:60],
            str(a.get("duration_days") or a.get("duration") or ""),
            str(a.get("early_start_day") if a.get("early_start_day") is not None else ""),
            str(a.get("early_finish_day") if a.get("early_finish_day") is not None else ""),
            str(a.get("total_float_days") if a.get("total_float_days") is not None else ""),
            "YES" if a.get("critical") else "",
        ])
    t = Table(data, repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    for r, row in enumerate(data):
        if r > 0 and row[7] == "YES":
            style.append(("BACKGROUND", (0, r), (-1, r), colors.HexColor("#F8CBAD")))
    t.setStyle(TableStyle(style))
    elements.append(t)
    if notes:
        elements += [Spacer(1, 12), Paragraph("<b>Notes</b>", styles["Normal"]),
                     Paragraph(notes, styles["Normal"])]
    pdf.build(elements)
    return path


@router.post("/v1/projects/{project_id}/export/schedule")
async def export_schedule(
    project_id: str,
    req: ScheduleExportRequest,
    format: str = "xlsx",
    auth: Dict[str, Any] = Depends(require_user),
):
    """Export an activity list to xlsx, docx, or pdf.

    `format` query param: `xlsx` (default), `docx`, or `pdf`.
    Body: `{"activities": [...], "project_name": "...", "start_date": "...", "notes": "..."}`.
    Returns the file via FileResponse — the browser triggers a download.
    """
    proj = _check_owner(project_id, auth["user_id"])
    if not req.activities:
        raise HTTPException(400, "activities is required and must be non-empty")
    name = req.project_name or proj.get("name") or "Schedule"

    fmt = (format or "xlsx").lower()
    if fmt == "xlsx":
        path = _render_xlsx(req.activities, name, req.start_date, req.notes)
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ext = "xlsx"
    elif fmt == "docx":
        path = _render_docx(req.activities, name, req.start_date, req.notes)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif fmt == "pdf":
        path = _render_pdf(req.activities, name, req.start_date, req.notes)
        media = "application/pdf"
        ext = "pdf"
    else:
        raise HTTPException(400, f"Unsupported format '{format}'. Use xlsx, docx, or pdf.")

    download_name = f"{name.replace(' ', '_')}_schedule.{ext}"
    return FileResponse(path, media_type=media, filename=download_name)


@router.post("/v1/projects/{project_id}/export/cost-boq")
async def export_cost_boq(
    project_id: str,
    req: CostBoqExportRequest,
    background_tasks: BackgroundTasks,
    auth: Dict[str, Any] = Depends(require_user),
):
    """Generate a FORMULA-LINKED cost-BOQ workbook (Cover / BOQ_Detail with
    =Qty*Rate / BOQ_Summary with cross-sheet links + % + cumulative /
    Cost_Charts). Pass ``categories`` directly, or ``document_id`` to derive
    them from an uploaded priced BOQ. Returns the .xlsx as a download AND (unless
    ``ingest`` is false) persists it into the project so it lands in RAG."""
    proj = _check_owner(project_id, auth["user_id"])
    name = req.project_name or proj.get("name") or "Project"
    categories = req.categories
    if not categories and req.document_id:
        categories = _categories_from_document(project_id, req.document_id)
    if not categories:
        raise HTTPException(400, "provide `categories` or a `document_id` to derive them from")
    from app.lib.boq_excel import generate_cost_boq
    meta = {
        "title": req.title or f"{name} — Bill of Quantities",
        "project": name, "location": req.location or "",
        "currency": req.currency, "date": req.date or "",
    }
    wb = generate_cost_boq(meta, categories)
    fd, path = tempfile.mkstemp(prefix="cost_boq_", suffix=".xlsx")
    os.close(fd)
    wb.save(path)

    # Persist the generated workbook back into the project as a document and
    # queue eager indexing — this is the link that gets the generated BOQ into
    # the project's RAG corpus so chat can answer from it. Mirrors the upload
    # path in app/routers/projects.py (add_document). Best-effort: a RAG-ingest
    # failure must NEVER 500 the export — the download below is the contract.
    if req.ingest:
        try:
            with open(path, "rb") as fh:
                raw_bytes = fh.read()
            # "boq" is not one of store.VALID_ROLES, so leave role unset and let
            # add_document classify it (falls back to "other").
            original_name = f"{name} - Cost BOQ (generated).xlsx"
            file_id = str(uuid.uuid4())[:8]
            stored_as = f"{file_id}_{original_name}"
            stored_path = os.path.join(DATA_DIR, stored_as)
            file_crypto.write_document(stored_path, raw_bytes)
            doc = projects_store.add_document(
                project_id, original_name, stored_as, stored_path,
                len(raw_bytes), role=None,
            )
            background_tasks.add_task(
                doc_index.maybe_eager_index, project_id, doc["id"])
        except Exception:  # noqa: BLE001 - ingest is best-effort, never fatal
            logger.warning(
                "cost-boq RAG ingest failed for project %s; download still served",
                project_id, exc_info=True)

    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=f"{name.replace(' ', '_')}_cost_BOQ.xlsx",
    )


def _persist_and_index(
    project_id: str, name: str, path: str, background_tasks: BackgroundTasks,
) -> None:
    """Persist a generated workbook back into the project and queue eager
    indexing so it lands in the project's RAG corpus. Best-effort: a RAG-ingest
    failure must NEVER break the download. Mirrors export_cost_boq's block."""
    try:
        with open(path, "rb") as fh:
            raw_bytes = fh.read()
        original_name = f"{name} - Priced BOQ (ESTIMATED from rate-card).xlsx"
        file_id = str(uuid.uuid4())[:8]
        stored_as = f"{file_id}_{original_name}"
        stored_path = os.path.join(DATA_DIR, stored_as)
        file_crypto.write_document(stored_path, raw_bytes)
        doc = projects_store.add_document(
            project_id, original_name, stored_as, stored_path,
            len(raw_bytes), role=None,
        )
        background_tasks.add_task(
            doc_index.maybe_eager_index, project_id, doc["id"])
    except Exception:  # noqa: BLE001 - ingest is best-effort, never fatal
        logger.warning(
            "price-boq RAG ingest failed for project %s; download still served",
            project_id, exc_info=True)


@router.post("/v1/projects/{project_id}/price-boq")
async def price_boq(
    project_id: str,
    req: PriceBoqRequest,
    background_tasks: BackgroundTasks,
    auth: Dict[str, Any] = Depends(require_user),
):
    """Price an UNPRICED BOQ from the typed rate-card and return a formula-linked
    cost workbook. The rates are ESTIMATED from the rate-card medians (NOT real
    tendered prices) -- the title/meta say so plainly. Persists + eager-indexes
    the workbook unless ``ingest`` is false. The download is the contract; a
    RAG-ingest failure never breaks it."""
    from app.lib import boq_pricing

    proj = _check_owner(project_id, auth["user_id"])
    name = req.project_name or proj.get("name") or "Project"

    # Validate asset_type / currency against the deployed rate-card.
    assets = boq_pricing.available_assets()
    if req.asset_type not in assets:
        raise HTTPException(
            400, f"Unknown asset_type '{req.asset_type}'. Valid options: "
                 f"{ {a: c for a, c in assets.items()} }")
    currencies = assets[req.asset_type]
    currency = req.currency or currencies[0]
    if currency not in currencies:
        raise HTTPException(
            400, f"Unknown currency '{currency}' for asset_type "
                 f"'{req.asset_type}'. Valid: {currencies}")

    # Extract UNPRICED line items from the document (rate is 0 -- expected).
    doc = projects_store.get_document(req.document_id)
    if not doc or not doc.get("file_path"):
        raise HTTPException(404, "document not found")
    from app.blocks.boq_processor import BOQProcessorBlock
    res = await BOQProcessorBlock().process(
        {"file_path": doc["file_path"], "project_id": project_id})
    if res.get("status") != "success" or not res.get("line_items"):
        raise HTTPException(
            422, "could not extract line items from that document -- a "
                 "digital/xlsx BOQ is required (scanned PDFs won't parse).")

    priced, summary = boq_pricing.price_line_items(
        res["line_items"], req.asset_type, currency)

    # Group priced items into cost-BOQ categories by their section when the
    # source BOQ carries one, else by the derived work_category. boq_processor
    # defaults section to "General" when there is no section column, so treat
    # "General" as absent and fall back to the work_category grouping.
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for it in priced:
        sec = it.get("section")
        key = sec if (sec and sec != "General") else it.get("work_category") or "Other"
        grouped.setdefault(key, []).append({
            "item_no": it["item_no"],
            "description": it["description"],
            "unit": it["unit"],
            "qty": it["qty"],
            "rate": it["rate"],
        })
    categories = [{"name": k, "items": v} for k, v in grouped.items()]

    from app.lib.boq_excel import generate_cost_boq
    meta = {
        "title": f"{name} - Cost BOQ (ESTIMATED from rate-card)",
        "project": name,
        "location": "",
        "currency": currency,
        "date": (
            f"Rates ESTIMATED from rate-card medians ({req.asset_type}/{currency}) "
            f"-- NOT tendered prices. Priced {summary['exact'] + summary['fallback']}"
            f"/{summary['total']} lines ({summary['no_rate']} flagged NO RATE)."
        ),
    }
    wb = generate_cost_boq(meta, categories)
    fd, path = tempfile.mkstemp(prefix="priced_boq_", suffix=".xlsx")
    os.close(fd)
    wb.save(path)

    if req.ingest:
        _persist_and_index(project_id, name, path, background_tasks)

    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=f"{name.replace(' ', '_')}_priced_BOQ_ESTIMATED.xlsx",
    )


_XLSX_MEDIA = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@router.post("/v1/projects/{project_id}/export/cost-schedule")
async def export_cost_schedule(
    project_id: str,
    req: CostScheduleExportRequest,
    auth: Dict[str, Any] = Depends(require_user),
):
    """Generate a cost-loaded L2 schedule workbook: CPM (ES/EF/LS/LF/float/
    critical) + cost per activity, a cumulative cost baseline (S-curve, live
    =prev+curr), a manpower histogram (man-days =Dur*Manpower), and a summary."""
    proj = _check_owner(project_id, auth["user_id"])
    if not req.activities:
        raise HTTPException(400, "activities is required and must be non-empty")
    name = req.project_name or proj.get("name") or "Project"
    from app.lib.pm_excel import generate_cost_loaded_schedule
    wb = generate_cost_loaded_schedule({"project": name, "currency": req.currency}, req.activities)
    fd, path = tempfile.mkstemp(prefix="cost_sched_", suffix=".xlsx"); os.close(fd)
    wb.save(path)
    return FileResponse(path, media_type=_XLSX_MEDIA,
                        filename=f"{name.replace(' ', '_')}_cost_loaded_schedule.xlsx")


@router.post("/v1/projects/{project_id}/export/evm")
async def export_evm(
    project_id: str,
    req: EvmExportRequest,
    auth: Dict[str, Any] = Depends(require_user),
):
    """Generate an EVM workbook from PV/EV/AC + BAC. CV/SV/CPI/SPI/EAC/ETC/VAC
    are all live formulas so the client can audit performance."""
    proj = _check_owner(project_id, auth["user_id"])
    if not req.periods:
        raise HTTPException(400, "periods is required and must be non-empty")
    name = req.project_name or proj.get("name") or "Project"
    from app.lib.pm_excel import generate_evm_workbook
    wb = generate_evm_workbook(
        {"project": name, "currency": req.currency, "bac": req.bac}, req.periods)
    fd, path = tempfile.mkstemp(prefix="evm_", suffix=".xlsx"); os.close(fd)
    wb.save(path)
    return FileResponse(path, media_type=_XLSX_MEDIA,
                        filename=f"{name.replace(' ', '_')}_EVM.xlsx")


# ── Conversation message export (Document Export Layer phase 1) ────────────
#
# Turn the latest assistant message in a conversation into a downloadable
# Word document. PDF + XLSX are 501 for now; they ride the same shape.

def _sanitize_filename(s: str) -> str:
    """Drop characters that filesystems / Content-Disposition headers dislike."""
    safe = "".join(c if c.isalnum() or c in "-_." else "_" for c in s)
    return safe.strip("_") or "export"


def _render_message_docx(
    project_name: str,
    assistant_text: str,
    conversation_id: str,
    message_index: int,
) -> str:
    """Render a single assistant message to a temp DOCX and return its path.

    The body keeps paragraph breaks but does not parse markdown -- bold/italic
    rendering is a follow-up. Markdown lives in the rendered text as-is, which
    Word will display literally. Sources are not yet persisted in agent_memory,
    so we emit a placeholder Sources section pointing to the live UI.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = doc.add_heading(f"{project_name} - Conversation Excerpt", level=0)
    for run in title.runs:
        run.font.size = Pt(18)

    subtitle = doc.add_paragraph(
        datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    )
    subtitle.runs[0].italic = True

    doc.add_paragraph()  # spacer

    # Body: split on blank lines so paragraph structure survives.
    for block in (assistant_text or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        doc.add_paragraph(block)

    doc.add_paragraph()
    doc.add_heading("Sources", level=2)
    doc.add_paragraph(
        "Sources are visible in the chat UI under each assistant message. "
        "Long-term source persistence to the message store is a future "
        "iteration; for now, refer to the live conversation at "
        "https://the-fork.onrender.com."
    )

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Generated by The Shovel - the-fork.onrender.com"
    )
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(8)

    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, prefix=f"export-{conversation_id[:8]}-"
    ) as f:
        path = f.name
    doc.save(path)
    return path


@router.post("/v1/projects/{project_id}/conversations/{conversation_id}/export")
async def export_conversation_message(
    project_id: str,
    conversation_id: str,
    format: str = Query("docx", pattern="^(docx|pdf|xlsx)$"),
    message_index: int = Query(
        -1,
        description="Index of the assistant message to export; -1 = most recent",
    ),
    auth: Dict[str, Any] = Depends(require_user),
):
    """Export one assistant message from a conversation as a downloadable file.

    Currently only ``format=docx`` is implemented; ``pdf`` and ``xlsx`` return
    501 so the frontend can hide / disable those format buttons cleanly.
    """
    proj = _check_owner(project_id, auth["user_id"])
    project_name = proj.get("name") or "Project"

    msgs = agent_memory.get_messages(conversation_id, limit=200)
    assistant_msgs = [m for m in msgs if m.get("role") == "assistant"]
    if not assistant_msgs:
        raise HTTPException(404, "No assistant messages in this conversation")

    # message_index is the position within assistant_msgs; -1 means newest.
    try:
        chosen = assistant_msgs[message_index]
    except IndexError:
        raise HTTPException(404, f"message_index {message_index} out of range")

    fmt = (format or "docx").lower()
    if fmt == "docx":
        path = _render_message_docx(
            project_name,
            chosen.get("content") or "",
            conversation_id,
            message_index,
        )
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif fmt in {"pdf", "xlsx"}:
        raise HTTPException(
            501,
            f"format='{fmt}' is not yet implemented; only 'docx' is available "
            "in this iteration of the Document Export Layer.",
        )
    else:
        raise HTTPException(400, f"Unsupported format '{format}'")

    safe_name = _sanitize_filename(project_name)
    download_name = f"{safe_name}-{conversation_id[:8]}.{ext}"
    return FileResponse(path, media_type=media, filename=download_name)
