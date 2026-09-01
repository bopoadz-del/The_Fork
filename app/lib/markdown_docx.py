"""Render an assistant message's markdown into a Word document properly.

THE INCIDENT (H1, gate battery ``13b2bf7``, 2026-08-31). The exported docx
was walked entry by entry through its ZIP central directory rather than
trusted by content-type, and ``word/document.xml`` contained, literally::

    The Accepted Contract Amount, **excluding VAT**, is:**SAR 1,754,504,456.25**

Asterisks as characters. The exporter's own docstring admitted it -- "does
not parse markdown ... Word will display literally" -- so this was a known
gap that had been shipping to clients. Same run found the document had no
``word/footer*.xml`` part at all: what the code called a footer was a
paragraph at the end of the BODY, which is not a footer and does not repeat
on the page.

Both are bugs against a stated requirement, which is why R3/R4 splits them
from the capability gaps: markdown rendering and the page footer land now;
multi-message report export (H1b) is a spec gap and does not.

WHAT IS SUPPORTED, and why the list stops where it does. Only the shapes
answers actually produce: ATX headings, bullet and numbered lists, pipe
tables, fenced code, blockquotes, horizontal rules, and inline ``**bold**``,
``*italic*``, ``` `code` ```.

``_underscore italics_`` is deliberately NOT supported. Construction answers
are full of ``snake_case`` identifiers -- ``message_index``, ``generate_wbs``,
``duration_overrides_applied`` -- and treating ``_`` as an emphasis delimiter
turns the middle of an identifier into italics and eats the underscores. A
renderer that corrupts a symbol name to gain italics has made the document
less true, not more finished.
"""

from __future__ import annotations

import re
from typing import Any, Iterable, List, Tuple

__all__ = [
    "render_markdown_into_docx",
    "strip_inline_markdown",
    "split_blocks",
]

# Inline spans, longest delimiter first so ** wins over *.
# Each alternative demands a non-space just inside the delimiters, so "2 * 3 *
# 4" and a bare asterisk are left alone.
_INLINE_RE = re.compile(
    r"(?P<code>`(?P<code_t>[^`\n]+?)`)"
    r"|(?P<bold>\*\*(?P<bold_t>[^\s*](?:[^*]*[^\s*])?)\*\*)"
    r"|(?P<ital>\*(?P<ital_t>[^\s*](?:[^*]*[^\s*])?)\*)"
)

_HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.+?)\s*#*$")
_BULLET_RE = re.compile(r"^\s{0,6}[-*•]\s+(?P<text>.+)$")
_NUMBER_RE = re.compile(r"^\s{0,6}\d+[.)]\s+(?P<text>.+)$")
_QUOTE_RE = re.compile(r"^\s{0,3}>\s?(?P<text>.*)$")
_HRULE_RE = re.compile(r"^\s{0,3}(?:-{3,}|\*{3,}|_{3,})\s*$")
_FENCE_RE = re.compile(r"^\s{0,3}(?:```|~~~)(?P<lang>[A-Za-z0-9+#-]*)\s*$")
_TABLE_SEP_CELL_RE = re.compile(r"^:?-{1,}:?$")


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _table_cells(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    if not _is_table_row(line):
        return False
    cells = _table_cells(line)
    return bool(cells) and all(_TABLE_SEP_CELL_RE.match(c.replace(" ", "")) for c in cells)


def iter_inline(text: str) -> Iterable[Tuple[str, str]]:
    """Yield ``(style, text)`` runs. ``style`` is "", "bold", "italic" or "code"."""
    pos = 0
    for m in _INLINE_RE.finditer(text or ""):
        if m.start() > pos:
            yield ("", text[pos:m.start()])
        if m.group("code"):
            yield ("code", m.group("code_t"))
        elif m.group("bold"):
            yield ("bold", m.group("bold_t"))
        else:
            yield ("italic", m.group("ital_t"))
        pos = m.end()
    if pos < len(text or ""):
        yield ("", text[pos:])


def strip_inline_markdown(text: str) -> str:
    """Plain text with the inline markers removed.

    Used where a target cannot carry mixed runs -- an Excel cell, a table
    header we style wholesale -- so the markers never reach a reader either
    way.
    """
    return "".join(t for _, t in iter_inline(text or ""))


def split_blocks(text: str) -> List[dict]:
    """Parse markdown into a flat list of block dicts.

    Kept separate from any docx call so the parse is testable without Word,
    and so the xlsx path can reuse it.
    """
    blocks: List[dict] = []
    para: List[str] = []

    def flush_para() -> None:
        if para:
            blocks.append({"kind": "para", "text": " ".join(para).strip()})
            para.clear()

    lines = (text or "").replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        fence = _FENCE_RE.match(line)
        if fence:
            flush_para()
            body: List[str] = []
            i += 1
            while i < len(lines) and not _FENCE_RE.match(lines[i]):
                body.append(lines[i])
                i += 1
            i += 1  # closing fence (or end of text)
            blocks.append({"kind": "code", "text": "\n".join(body),
                           "lang": fence.group("lang") or ""})
            continue

        if _is_table_row(line):
            flush_para()
            rows: List[List[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                if not _is_table_separator(lines[i]):
                    rows.append(_table_cells(lines[i]))
                i += 1
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            continue

        if not line.strip():
            flush_para()
            i += 1
            continue

        if _HRULE_RE.match(line):
            flush_para()
            blocks.append({"kind": "rule"})
            i += 1
            continue

        h = _HEADING_RE.match(line)
        if h:
            flush_para()
            blocks.append({"kind": "heading", "level": len(h.group("hashes")),
                           "text": h.group("text")})
            i += 1
            continue

        q = _QUOTE_RE.match(line)
        if q:
            flush_para()
            blocks.append({"kind": "quote", "text": q.group("text")})
            i += 1
            continue

        b = _BULLET_RE.match(line)
        if b:
            flush_para()
            blocks.append({"kind": "bullet", "text": b.group("text")})
            i += 1
            continue

        n = _NUMBER_RE.match(line)
        if n:
            flush_para()
            blocks.append({"kind": "number", "text": n.group("text")})
            i += 1
            continue

        para.append(line.strip())
        i += 1

    flush_para()
    return blocks


def _add_runs(paragraph: Any, text: str) -> None:
    for style, chunk in iter_inline(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if style == "bold":
            run.bold = True
        elif style == "italic":
            run.italic = True
        elif style == "code":
            run.font.name = "Consolas"


def _styled_paragraph(doc: Any, style: str | None):
    """add_paragraph(style=...) raises KeyError on a template without that
    style. Falling back to an unstyled paragraph keeps the export working on
    any base template rather than 500-ing on a missing list style."""
    if style:
        try:
            return doc.add_paragraph(style=style)
        except KeyError:
            pass
    return doc.add_paragraph()


def render_markdown_into_docx(doc: Any, text: str) -> int:
    """Append ``text``'s markdown to ``doc``. Returns the block count."""
    from docx.shared import Pt

    blocks = split_blocks(text)
    for blk in blocks:
        kind = blk["kind"]
        if kind == "heading":
            # Word's built-in headings stop at 9; answers never go past 6.
            doc.add_heading(strip_inline_markdown(blk["text"]),
                            level=min(max(blk["level"], 1), 6))
        elif kind == "bullet":
            _add_runs(_styled_paragraph(doc, "List Bullet"), blk["text"])
        elif kind == "number":
            _add_runs(_styled_paragraph(doc, "List Number"), blk["text"])
        elif kind == "quote":
            p = _styled_paragraph(doc, "Quote")
            _add_runs(p, blk["text"])
            for r in p.runs:
                r.italic = True
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(blk["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "rule":
            doc.add_paragraph()
        elif kind == "table":
            rows = blk["rows"]
            width = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=width)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for r_i, row in enumerate(rows):
                for c_i in range(width):
                    cell = table.cell(r_i, c_i)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    _add_runs(para, row[c_i] if c_i < len(row) else "")
                    if r_i == 0:
                        for run in para.runs:
                            run.bold = True
        else:
            _add_runs(doc.add_paragraph(), blk["text"])
    return len(blocks)
