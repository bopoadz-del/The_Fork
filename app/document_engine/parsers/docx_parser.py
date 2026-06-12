"""DOCX Parser — Layer 1 syntactic extraction for document engine."""
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from pathlib import Path


@dataclass
class DOCXDocument:
    source: str
    text: str = ""
    paragraphs: List[str] = field(default_factory=list)
    headings: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    lists: List[str] = field(default_factory=list)


class DOCXParser:
    """Extract sections, headings, bullets, requirement tables from Word."""

    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}

    def parse(self, file_path: str) -> DOCXDocument:
        path = Path(file_path)
        doc = DOCXDocument(source=str(path))

        try:
            from docx import Document
            document = Document(file_path)

            for para in document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    level = 0
                    try:
                        level = int(style.replace("Heading ", ""))
                    except ValueError:
                        pass
                    doc.headings.append({"level": level, "text": text})
                elif para.style and "List" in style:
                    doc.lists.append(text)
                else:
                    doc.paragraphs.append(text)

                doc.text += text + "\n"

            for table in document.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                doc.tables.append(rows)

        except ImportError:
            # Fallback — raw text read (garbled but may yield something)
            doc.text = path.read_text(errors="ignore")
            doc.paragraphs = [p for p in doc.text.split("\n") if p.strip()]

        return doc
